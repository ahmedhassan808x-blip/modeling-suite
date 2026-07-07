"""
Memo builder — Word documents in the same navy & steel identity as the
decks: navy headings, banded tables, embedded theme charts.
"""

from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from shared.theme import FONT, GRAY, INK, MIST, NAVY, SLATE

C_NAVY, C_INK = RGBColor.from_string(NAVY), RGBColor.from_string(INK)
C_SLATE, C_GRAY = RGBColor.from_string(SLATE), RGBColor.from_string(GRAY)


def _shade(cell, hex_color):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(el)


class Memo:
    def __init__(self, title, subtitle):
        self.doc = Document()
        style = self.doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10.5)
        style.font.color.rgb = C_INK

        t = self.doc.add_paragraph()
        r = t.add_run(title)
        r.font.size, r.font.bold, r.font.color.rgb = Pt(20), True, C_NAVY
        s = self.doc.add_paragraph()
        r = s.add_run(f"{subtitle}  |  {date.today().strftime('%B %d, %Y')}")
        r.font.size, r.font.color.rgb = Pt(10), C_SLATE

    def heading(self, text, level=1):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(14 if level == 1 else 11.5)
        r.font.color.rgb = C_NAVY if level == 1 else C_SLATE
        p.space_before = Pt(12)
        return p

    def para(self, text, italic=False, small=False):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.font.italic = italic
        if small:
            r.font.size, r.font.color.rgb = Pt(8.5), C_GRAY
        return p

    def bullets(self, items):
        for it in items:
            self.doc.add_paragraph(it, style="List Bullet")

    def table(self, rows, col_widths=None):
        tbl = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        tbl.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = tbl.cell(i, j)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.text = ""
                p = cell.paragraphs[0]
                if j > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r = p.add_run(str(val))
                r.font.size = Pt(9)
                r.font.name = FONT
                if i == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor.from_string("FFFFFF")
                    _shade(cell, NAVY)
                elif i % 2 == 0:
                    _shade(cell, MIST)
                if col_widths and i == 0:
                    tbl.columns[j].width = Inches(col_widths[j])
        return tbl

    def image(self, png, width_in=6.3):
        self.doc.add_picture(str(png), width=Inches(width_in))

    def save(self, path):
        self.doc.save(str(path))
        return path
